# shared.py — Umumiy funksiyalar, konstantalar, DB va CSS
import sqlite3
import hashlib
import io
from datetime import datetime
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATABASE = "adliya_murojaat.db"

KATEGORIYALAR = [
    "Tadbirkorlik", "Transport va yo'llar", "Uy-joy va qurilish", "Ta'lim",
    "Sog'liqni saqlash", "Soliq va moliya", "Mehnat munosabatlari",
    "Yer masalalari", "Ekologiya", "Boshqa",
]
KATEGORIYALAR_RU = [
    "Предпринимательство", "Транспорт и дороги", "Жилье и строительство", "Образование",
    "Здравоохранение", "Налоги и финансы", "Трудовые отношения",
    "Земельные вопросы", "Экология", "Другое",
]
KATEGORIYALAR_EN = [
    "Entrepreneurship", "Transport and Roads", "Housing and Construction", "Education",
    "Healthcare", "Tax and Finance", "Labor Relations",
    "Land Issues", "Ecology", "Other",
]

STATUS_NOMLARI = {
    "yangi": "🆕 Yangi",
    "ko_rib_chiqilmoqda": "🟡 Ko'rib chiqilmoqda",
    "asosli": "✅ Asosli deb topildi",
    "asossiz": "❌ Asossiz deb topildi",
    "loyiha_tayyorlandi": "📄 Loyiha tayyorlandi",
    "kiritildi": "🚀 Yuqoriga kiritildi",
}
STATUS_NOMLARI_RU = {
    "yangi": "🆕 Новый",
    "ko_rib_chiqilmoqda": "🟡 На рассмотрении",
    "asosli": "✅ Признано обоснованным",
    "asossiz": "❌ Признано необоснованным",
    "loyiha_tayyorlandi": "📄 Проект подготовлен",
    "kiritildi": "🚀 Внесено наверх",
}
STATUS_NOMLARI_EN = {
    "yangi": "🆕 New",
    "ko_rib_chiqilmoqda": "🟡 Under Review",
    "asosli": "✅ Found Substantiated",
    "asossiz": "❌ Found Unsubstantiated",
    "loyiha_tayyorlandi": "📄 Draft Prepared",
    "kiritildi": "🚀 Submitted Upward",
}

STATUS_TARTIBI = list(STATUS_NOMLARI.keys())
STATUS_RANGLAR = {
    "yangi": "#64748B", "ko_rib_chiqilmoqda": "#F59E0B",
    "asosli": "#16A34A", "asossiz": "#DC2626",
    "loyiha_tayyorlandi": "#0EA5E9", "kiritildi": "#2563EB",
}

# ============================================================
# TIL FUNKSIYALARI
# ============================================================
LANG = {
    "uz": {
        "name": "O'zbek",
        "header_title": "O'ZBEKISTON RESPUBLIKASI ADLIYA VAZIRLIGI",
        "header_subtitle": "Fuqarolar tashabbusi platformasi",
        "nav_home": "🏠 Bosh sahifa",
        "nav_submit": "✍️ Taklif yuborish",
        "nav_reestr": "📋 Ochiq reestr",
        "nav_login": "🔐 Xodim paneli",
        "nav_appeals": "📂 Murojaatlar",
        "nav_stats": "📊 Statistika",
        "nav_logout": "🚪 Chiqish",
        "hero_title": "⚖️ Qonunchilikni birgalikda takomillashtiramiz",
        "hero_desc": "Amaldagi qonun yoki tartibda kamchilik ko'rdingizmi? O'z taklifingizni bildiring. Biz o'rganamiz, tahlil qilamiz va asosli bo'lsa — o'zgartirish kiritish uchun harakat qilamiz.",
        "hero_btn": "✍️ Taklif yuborish",
        "stats_title": "📊 Platforma statistikasi",
        "stats_total": "Jami murojaatlar",
        "stats_valid": "Asosli deb topilgan",
        "stats_draft": "Loyiha tayyorlangan",
        "stats_submitted": "Yuqoriga kiritilgan",
        "how_title": "📋 Qanday ishlaydi?",
        "step1_title": "Muammoni toping",
        "step1_desc": "Kundalik hayotda amaldagi qonun yoki tartibdagi kamchilikni aniqlang",
        "step2_title": "Bizga yozing",
        "step2_desc": "Muammoni, asosini va taklifingizni platformaga yuboring",
        "step3_title": "Biz o'rganamiz",
        "step3_desc": "Mutaxassislarimiz holatni tahlil qiladi, statistika yig'adi",
        "step4_title": "Natija",
        "step4_desc": "Asosli bo'lsa — o'zgartirish loyihasi tayyorlanadi va kiritiladi",
        "submit_title": "✍️ Tashabbus bilan murojaat qilish",
        "submit_desc": "Amaldagi qonun yoki tartibda kamchilik ko'rdingizmi? Quyidagi formani to'ldiring. Sizning fikringiz muhim — asosli bo'lsa, biz qonunchilikni o'zgartirish bo'yicha harakat qilamiz.",
        "personal_info": "👤 Shaxsiy ma'lumotlar",
        "full_name": "Ism-familiyangiz *",
        "phone": "Telefon raqamingiz *",
        "email_opt": "Email (ixtiyoriy)",
        "problem_title": "📝 Muammo haqida",
        "category": "Kategoriya *",
        "category_placeholder": "Tanlang...",
        "problem": "Muammo nima? *",
        "law_basis": "Qaysi qonun/tartib asosida? (ixtiyoriy)",
        "why_wrong": "Nima uchun bu xato yoki adolatsiz? *",
        "suggestion": "Sizningcha qanday bo'lishi kerak? *",
        "required_note": "* bilan belgilangan maydonlar majburiy",
        "submit_btn": "📤 Yuborish",
        "clear_btn": "🔄 Tozalash",
        "success_msg": "✅ Murojaatingiz qabul qilindi!",
        "your_number": "Sizning murojaat raqamingiz",
        "save_number": "Bu raqamni saqlab qo'ying. Ochiq reestr bo'limida murojaatingizni kuzatishingiz mumkin.",
        "required_error": "Iltimos, quyidagi majburiy maydonlarni to'ldiring",
        "reestr_title": "📋 Ochiq reestr",
        "reestr_desc": "Barcha murojaatlarning ochiq ma'lumotlari (fuqaro shaxsiy ma'lumotlarisiz)",
        "search": "🔍 Qidirish",
        "search_placeholder": "Kalit so'zni kiriting...",
        "status_filter": "Status",
        "category_filter": "Kategoriya",
        "all": "Barchasi",
        "total_found": "ta murojaat topildi",
        "not_found": "Murojaatlar topilmadi",
        "detail": "📄 Batafsil",
        "back": "← Orqaga",
        "view": "👁 Ko'rish",
        "delete": "🗑 O'chirish",
        "content": "📝 Murojaat mazmuni",
        "response": "📩 Xodim javobi",
        "reviewed_at": "Ko'rib chiqilgan",
        "not_indicated": "Ko'rsatilmagan",
        "login_title": "🔐 Xodim paneliga kirish",
        "login_username": "Login",
        "login_password": "Parol",
        "login_btn": "Kirish",
        "login_default": "Default: login: **admin**, parol: **admin123**",
        "login_error": "Login yoki parol xato",
        "login_success": "Tizimga muvaffaqiyatli kirdingiz",
        "appeals_title": "📂 Murojaatlar ro'yxati",
        "new": "Yangi",
        "in_progress": "Jarayonda",
        "valid": "Asosli",
        "filter_status": "Status bo'yicha filtrlash",
        "stats_page_title": "📈 Statistika",
        "stats_by_category": "📊 Kategoriyalar bo'yicha",
        "stats_by_status": "📈 Statuslar bo'yicha",
        "no_data": "Ma'lumot yo'q",
        "decision_title": "⚖️ Qaror qabul qilish",
        "change_status": "Statusni o'zgartirish *",
        "internal_note": "Ichki izoh (faqat xodimlar ko'radi)",
        "response_to_citizen": "Fuqaroga javob (fuqaro ko'radi) *",
        "save_btn": "💾 Saqlash",
        "response_required": "Fuqaroga javobni kiriting.",
        "updated_success": "✅ Murojaat muvaffaqiyatli yangilandi",
        "deleted_success": "✅ Murojaat o'chirildi",
        "warning_xodim": "⚠️ Xodim sifatida taklif yubora olmaysiz. Iltimos, avval tizimdan chiqing.",
        "footer_address": "100000, Toshkent shahri, Mustaqillik maydoni, 5-uy",
        "footer_contact": "Tel: +998 (71) 200-00-00 | E-mail: info@adliya.uz",
        "footer_about": "Platforma haqida",
        "footer_terms": "Foydalanish shartlari",
        "footer_privacy": "Maxfiylik siyosati",
        "footer_contact_link": "Aloqa",
        "footer_faq": "Ko'p beriladigan savollar",
        "footer_copyright": "O'zbekiston Respublikasi Adliya vazirligi. Barcha huquqlar himoyalangan.",
        "citizen_info": "👤 Fuqaro ma'lumotlari",
        "category_label": "Kategoriya",
        "date_sent": "Yuborilgan sana",
        "law_basis_label": "Qonun/tartib asosi",
        "why_wrong_label": "Nima uchun xato/adolatsiz",
        "suggestion_label": "Fuqaro taklifi",
        "wait": "Iltimos, avval tizimga kiring",
    },
    "ru": {
        "name": "Русский",
        "header_title": "МИНИСТЕРСТВО ЮСТИЦИИ РЕСПУБЛИКИ УЗБЕКИСТАН",
        "header_subtitle": "Платформа гражданских инициатив",
        "nav_home": "🏠 Главная",
        "nav_submit": "✍️ Отправить предложение",
        "nav_reestr": "📋 Открытый реестр",
        "nav_login": "🔐 Панель сотрудника",
        "nav_appeals": "📂 Обращения",
        "nav_stats": "📊 Статистика",
        "nav_logout": "🚪 Выход",
        "hero_title": "⚖️ Совершенствуем законодательство вместе",
        "hero_desc": "Заметили недостаток в действующем законе или порядке? Сообщите нам. Мы изучаем, анализируем и, если обосновано, предпринимаем действия для внесения изменений.",
        "hero_btn": "✍️ Отправить предложение",
        "stats_title": "📊 Статистика платформы",
        "stats_total": "Всего обращений",
        "stats_valid": "Признано обоснованными",
        "stats_draft": "Подготовлено проектов",
        "stats_submitted": "Внесено наверх",
        "how_title": "📋 Как это работает?",
        "step1_title": "Найдите проблему",
        "step1_desc": "Определите недостаток в действующем законе или порядке в повседневной жизни",
        "step2_title": "Напишите нам",
        "step2_desc": "Отправьте проблему, основание и ваше предложение на платформу",
        "step3_title": "Мы изучаем",
        "step3_desc": "Наши специалисты анализируют ситуацию, собирают статистику",
        "step4_title": "Результат",
        "step4_desc": "Если обосновано — подготавливается проект изменений и вносится",
        "submit_title": "✍️ Обратиться с инициативой",
        "submit_desc": "Заметили недостаток в действующем законе или порядке? Заполните форму ниже. Ваше мнение важно — если обосновано, мы предпримем действия по изменению законодательства.",
        "personal_info": "👤 Личные данные",
        "full_name": "ФИО *",
        "phone": "Номер телефона *",
        "email_opt": "Email (необязательно)",
        "problem_title": "📝 О проблеме",
        "category": "Категория *",
        "category_placeholder": "Выберите...",
        "problem": "В чем проблема? *",
        "law_basis": "На основании какого закона/порядка? (необязательно)",
        "why_wrong": "Почему это неверно или несправедливо? *",
        "suggestion": "Как должно быть по вашему мнению? *",
        "required_note": "* обязательные поля",
        "submit_btn": "📤 Отправить",
        "clear_btn": "🔄 Очистить",
        "success_msg": "✅ Ваше обращение принято!",
        "your_number": "Номер вашего обращения",
        "save_number": "Сохраните этот номер. Вы можете отслеживать статус обращения в разделе «Открытый реестр».",
        "required_error": "Пожалуйста, заполните следующие обязательные поля",
        "reestr_title": "📋 Открытый реестр",
        "reestr_desc": "Открытая информация по всем обращениям (без личных данных граждан)",
        "search": "🔍 Поиск",
        "search_placeholder": "Введите ключевое слово...",
        "status_filter": "Статус",
        "category_filter": "Категория",
        "all": "Все",
        "total_found": "обращений найдено",
        "not_found": "Обращения не найдены",
        "detail": "📄 Подробнее",
        "back": "← Назад",
        "view": "👁 Просмотр",
        "delete": "🗑 Удалить",
        "content": "📝 Содержание обращения",
        "response": "📩 Ответ сотрудника",
        "reviewed_at": "Рассмотрено",
        "not_indicated": "Не указано",
        "login_title": "🔐 Вход в панель сотрудника",
        "login_username": "Логин",
        "login_password": "Пароль",
        "login_btn": "Войти",
        "login_default": "По умолчанию: логин: **admin**, пароль: **admin123**",
        "login_error": "Неверный логин или пароль",
        "login_success": "Вы успешно вошли в систему",
        "appeals_title": "📂 Список обращений",
        "new": "Новые",
        "in_progress": "В процессе",
        "valid": "Обоснованные",
        "filter_status": "Фильтр по статусу",
        "stats_page_title": "📈 Статистика",
        "stats_by_category": "📊 По категориям",
        "stats_by_status": "📈 По статусам",
        "no_data": "Нет данных",
        "decision_title": "⚖️ Принятие решения",
        "change_status": "Изменить статус *",
        "internal_note": "Внутренняя заметка (видна только сотрудникам)",
        "response_to_citizen": "Ответ гражданину (виден гражданину) *",
        "save_btn": "💾 Сохранить",
        "response_required": "Введите ответ гражданину.",
        "updated_success": "✅ Обращение успешно обновлено",
        "deleted_success": "✅ Обращение удалено",
        "warning_xodim": "⚠️ Вы не можете отправлять предложения как сотрудник. Пожалуйста, выйдите из системы.",
        "footer_address": "100000, г. Ташкент, площадь Мустакиллик, 5",
        "footer_contact": "Тел: +998 (71) 200-00-00 | E-mail: info@adliya.uz",
        "footer_about": "О платформе",
        "footer_terms": "Условия использования",
        "footer_privacy": "Политика конфиденциальности",
        "footer_contact_link": "Контакты",
        "footer_faq": "Часто задаваемые вопросы",
        "footer_copyright": "Министерство юстиции Республики Узбекистан. Все права защищены.",
        "citizen_info": "👤 Данные гражданина",
        "category_label": "Категория",
        "date_sent": "Дата отправки",
        "law_basis_label": "Основание (закон/порядок)",
        "why_wrong_label": "Почему это неверно/несправедливо",
        "suggestion_label": "Предложение гражданина",
        "wait": "Пожалуйста, сначала войдите в систему",
    },
    "en": {
        "name": "English",
        "header_title": "MINISTRY OF JUSTICE OF THE REPUBLIC OF UZBEKISTAN",
        "header_subtitle": "Citizen Initiative Platform",
        "nav_home": "🏠 Home",
        "nav_submit": "✍️ Submit Proposal",
        "nav_reestr": "📋 Public Registry",
        "nav_login": "🔐 Staff Panel",
        "nav_appeals": "📂 Appeals",
        "nav_stats": "📊 Statistics",
        "nav_logout": "🚪 Logout",
        "hero_title": "⚖️ Improving Legislation Together",
        "hero_desc": "Have you noticed a flaw in current laws or procedures? Let us know. We review, analyze, and if justified — take action to make changes.",
        "hero_btn": "✍️ Submit Proposal",
        "stats_title": "📊 Platform Statistics",
        "stats_total": "Total Appeals",
        "stats_valid": "Found Substantiated",
        "stats_draft": "Drafts Prepared",
        "stats_submitted": "Submitted Upward",
        "how_title": "📋 How It Works?",
        "step1_title": "Identify the Problem",
        "step1_desc": "Find a flaw in current laws or procedures in everyday life",
        "step2_title": "Write to Us",
        "step2_desc": "Submit the problem, its basis, and your suggestion to the platform",
        "step3_title": "We Review",
        "step3_desc": "Our experts analyze the situation and collect statistics",
        "step4_title": "Result",
        "step4_desc": "If substantiated — a draft amendment is prepared and submitted",
        "submit_title": "✍️ Submit an Initiative",
        "submit_desc": "Have you noticed a flaw in current laws or procedures? Fill out the form below. Your opinion matters — if substantiated, we will take action to amend the legislation.",
        "personal_info": "👤 Personal Information",
        "full_name": "Full Name *",
        "phone": "Phone Number *",
        "email_opt": "Email (optional)",
        "problem_title": "📝 About the Problem",
        "category": "Category *",
        "category_placeholder": "Select...",
        "problem": "What is the problem? *",
        "law_basis": "Based on which law/procedure? (optional)",
        "why_wrong": "Why is this wrong or unfair? *",
        "suggestion": "How should it be in your opinion? *",
        "required_note": "* required fields",
        "submit_btn": "📤 Submit",
        "clear_btn": "🔄 Clear",
        "success_msg": "✅ Your appeal has been accepted!",
        "your_number": "Your appeal number",
        "save_number": "Save this number. You can track your appeal status in the Public Registry section.",
        "required_error": "Please fill in the following required fields",
        "reestr_title": "📋 Public Registry",
        "reestr_desc": "Public information on all appeals (without personal data of citizens)",
        "search": "🔍 Search",
        "search_placeholder": "Enter keyword...",
        "status_filter": "Status",
        "category_filter": "Category",
        "all": "All",
        "total_found": "appeal(s) found",
        "not_found": "No appeals found",
        "detail": "📄 Details",
        "back": "← Back",
        "view": "👁 View",
        "delete": "🗑 Delete",
        "content": "📝 Appeal Content",
        "response": "📩 Staff Response",
        "reviewed_at": "Reviewed at",
        "not_indicated": "Not indicated",
        "login_title": "🔐 Staff Panel Login",
        "login_username": "Username",
        "login_password": "Password",
        "login_btn": "Login",
        "login_default": "Default: login: **admin**, password: **admin123**",
        "login_error": "Invalid username or password",
        "login_success": "Successfully logged in",
        "appeals_title": "📂 Appeals List",
        "new": "New",
        "in_progress": "In Progress",
        "valid": "Substantiated",
        "filter_status": "Filter by Status",
        "stats_page_title": "📈 Statistics",
        "stats_by_category": "📊 By Category",
        "stats_by_status": "📈 By Status",
        "no_data": "No data",
        "decision_title": "⚖️ Decision Making",
        "change_status": "Change Status *",
        "internal_note": "Internal note (staff only)",
        "response_to_citizen": "Response to Citizen (citizen can see) *",
        "save_btn": "💾 Save",
        "response_required": "Please enter a response to the citizen.",
        "updated_success": "✅ Appeal successfully updated",
        "deleted_success": "✅ Appeal deleted",
        "warning_xodim": "⚠️ You cannot submit proposals as a staff member. Please log out first.",
        "footer_address": "100000, Tashkent, Mustaqillik Square, 5",
        "footer_contact": "Tel: +998 (71) 200-00-00 | E-mail: info@adliya.uz",
        "footer_about": "About Platform",
        "footer_terms": "Terms of Use",
        "footer_privacy": "Privacy Policy",
        "footer_contact_link": "Contact",
        "footer_faq": "FAQ",
        "footer_copyright": "Ministry of Justice of the Republic of Uzbekistan. All rights reserved.",
        "citizen_info": "👤 Citizen Information",
        "category_label": "Category",
        "date_sent": "Date Sent",
        "law_basis_label": "Law/Procedure Basis",
        "why_wrong_label": "Why Wrong/Unfair",
        "suggestion_label": "Citizen's Suggestion",
        "wait": "Please log in first",
    }
}


def t(key):
    """Berilgan kalit bo'yicha joriy tildagi matnni qaytaradi."""
    lang = st.session_state.get("til", "uz")
    return LANG[lang].get(key, LANG["uz"].get(key, key))


def get_kategoriyalar():
    lang = st.session_state.get("til", "uz")
    if lang == "ru": return KATEGORIYALAR_RU
    if lang == "en": return KATEGORIYALAR_EN
    return KATEGORIYALAR


def get_status_nomlari():
    lang = st.session_state.get("til", "uz")
    if lang == "ru": return STATUS_NOMLARI_RU
    if lang == "en": return STATUS_NOMLARI_EN
    return STATUS_NOMLARI


# ============================================================
# MA'LUMOTLAR BAZASI
# ============================================================
def get_conn():
    conn = sqlite3.connect(DATABASE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_conn()
    conn.execute("""CREATE TABLE IF NOT EXISTS xodimlar
                    (
                        id
                        INTEGER
                        PRIMARY
                        KEY
                        AUTOINCREMENT,
                        login
                        TEXT
                        UNIQUE
                        NOT
                        NULL,
                        parol_hash
                        TEXT
                        NOT
                        NULL,
                        ism
                        TEXT
                        NOT
                        NULL,
                        lavozim
                        TEXT,
                        yaratilgan
                        TEXT
                        DEFAULT
                        CURRENT_TIMESTAMP
                    )""")
    conn.execute("""CREATE TABLE IF NOT EXISTS murojaatlar
    (
        id
        INTEGER
        PRIMARY
        KEY
        AUTOINCREMENT,
        fuqaro_ism
        TEXT
        NOT
        NULL,
        telefon
        TEXT
        NOT
        NULL,
        email
        TEXT,
        kategoriya
        TEXT
        NOT
        NULL,
        muammo
        TEXT
        NOT
        NULL,
        qonun_asosi
        TEXT,
        nima_uchun_xato
        TEXT
        NOT
        NULL,
        taklif
        TEXT
        NOT
        NULL,
        status
        TEXT
        DEFAULT
        'yangi',
        xodim_id
        INTEGER,
        xodim_izohi
        TEXT,
        javob
        TEXT,
        yaratilgan
        TEXT
        DEFAULT
        CURRENT_TIMESTAMP,
        ko_rib_chiqilgan
        TEXT,
        FOREIGN
        KEY
                    (
        xodim_id
                    ) REFERENCES xodimlar
                    (
                        id
                    ))""")
    try:
        conn.execute("INSERT INTO xodimlar (login, parol_hash, ism, lavozim) VALUES (?, ?, ?, ?)",
                     ("admin", hashlib.sha256("admin123".encode()).hexdigest(), "Administrator", "Bosh mutaxassis"))
        conn.commit()
    except sqlite3.IntegrityError:
        pass
    conn.close()


def hash_parol(parol): return hashlib.sha256(parol.encode()).hexdigest()


# ============================================================
# SESSION STATE BOSHLANG'ICH QIYMATLARI
# ============================================================
def init_session_state():
    for k, v in [("sahifa", "Bosh sahifa"), ("xodim_id", None), ("xodim_ism", None), ("tanlangan_murojaat", None),
                 ("form_ism", ""), ("form_telefon", ""), ("form_email", ""), ("form_kategoriya", "Tanlang..."),
                 ("form_muammo", ""), ("form_qonun", ""), ("form_nima", ""), ("form_taklif", ""),
                 ("til", "uz")]:
        if k not in st.session_state: st.session_state[k] = v


# ============================================================
# TO'LIQ EKRAN + YANGI RANG PALITRASI (Primary/Secondary/Accent) CSS
# ============================================================
CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
* { font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif; }

/* ===== RANG O'ZGARUVCHILARI =====
Primary:   #1E3A8A
Secondary: #2563EB
Accent:    #F59E0B
Background:#F8FAFC
Surface:   #FFFFFF
Text:      #0F172A
Success:   #16A34A
Warning:   #F59E0B
Error:     #DC2626
Info:      #0EA5E9
*/

/* ===== TO'LIQ EKRAN ===== */
.stApp { background-color: #F8FAFC; margin: 0; padding: 0; }
.main .block-container { max-width: 100% !important; padding: 0 !important; margin: 0 !important; }
.main > div:first-child { padding: 0 !important; max-width: 100% !important; }
.main > div:first-child > div { width: 100% !important; max-width: 100% !important; padding: 0 !important; }
.element-container { max-width: 100% !important; }

/* ===== STREAMLIT XIZMAT ELEMENTLARINI YASHIRISH (Deploy, menyu, header) ===== */
.stApp > header { display: none !important; }
header[data-testid="stHeader"] { display: none !important; visibility: hidden !important; height: 0 !important; }
[data-testid="stToolbar"] { display: none !important; visibility: hidden !important; }
[data-testid="stDecoration"] { display: none !important; }
[data-testid="stStatusWidget"] { display: none !important; }
.stDeployButton { display: none !important; }
#MainMenu { visibility: hidden !important; display: none !important; }
footer { display: none !important; }
.st-emotion-cache-1jicfl2 { padding: 0 !important; }

/* ===== PAGE CONTAINER ===== */
.page-container { max-width: 1400px; margin: 0 auto; padding: 0 40px; }

/* ===== GOV HEADER (Primary -> Secondary gradient) ===== */
.gov-header {
    background: linear-gradient(180deg, #1E3A8A 0%, #1D4ED8 100%);
    padding: 20px 0 14px 0;
    color: white;
    width: 100%;
}
.gov-header .header-inner {
    max-width: 1400px;
    margin: 0 auto;
    padding: 0 40px;
    display: flex;
    align-items: center;
    justify-content: flex-start;
}
.gov-header .header-left { display: flex; align-items: center; gap: 18px; }
.gov-header .header-left .gerb {
    width: 62px; height: 62px;
    border-radius: 50%;
    object-fit: contain;
    background: rgba(255,255,255,0.9);
    padding: 4px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.15);
}
.gov-header .title-block h1 {
    font-size: 19px; font-weight: 800; margin: 0; line-height: 1.35;
    letter-spacing: 0.3px; text-transform: uppercase;
    text-shadow: 0 1px 2px rgba(0,0,0,0.12);
}

/* ===== TIL QATORI (sarlavhaning ostida, to'q ko'k fonda uzluksiz davom etadi) ===== */
div:has(> div.gov-header),
[data-testid="element-container"]:has(.gov-header) {
    margin-bottom: 0 !important;
}
div:has(> div.gov-header) + div,
[data-testid="element-container"]:has(.gov-header) + div {
    background: #1D4ED8 !important;
    margin-top: 0 !important;
    padding-top: 0 !important;
    border-bottom: 3px solid #F59E0B;
}
.header-lang-row div.stButton > button {
    background: rgba(255,255,255,0.12) !important;
    color: #ffffff !important;
    border: 1px solid rgba(255,255,255,0.5) !important;
    border-radius: 6px !important;
    padding: 4px 14px !important;
    font-size: 12.5px !important;
    font-weight: 700 !important;
    letter-spacing: 0.5px;
}
.header-lang-row div.stButton > button:hover {
    background: rgba(255,255,255,0.28) !important;
    border-color: #ffffff !important;
}
.header-lang-row div.stButton > button[kind="primary"] {
    background: #ffffff !important;
    color: #1E3A8A !important;
    border: 1px solid #ffffff !important;
}

/* ===== GOV NAV ===== */
.gov-nav {
    background: #FFFFFF;
    border-bottom: 1px solid #e0e4e8;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
    width: 100%;
    position: sticky;
    top: 0;
    z-index: 999;
    padding: 4px 0;
}
.gov-nav .nav-inner {
    max-width: 1400px;
    margin: 0 auto;
    display: flex;
    align-items: center;
    gap: 2px;
    padding: 0 40px;
}

/* ===== HERO ===== */
.gov-hero {
    background: linear-gradient(135deg, #1E3A8A 0%, #2563EB 50%, #3B82F6 100%);
    color: white; padding: 50px 50px; border-radius: 14px;
    margin: 25px 0 30px 0; position: relative; overflow: hidden;
    text-align: center;
}
.gov-hero::before {
    content: '🇺🇿'; position: absolute; top: -40px; right: -20px;
    font-size: 160px; opacity: 0.07;
}
.gov-hero h2 { font-size: 2em; font-weight: 700; margin: 0 0 12px 0; position: relative; line-height: 1.3; }
.gov-hero p { font-size: 1em; opacity: 0.9; max-width: 700px; margin: 0 auto 22px; line-height: 1.7; position: relative; }
.gov-hero .hero-btn {
    background: #F59E0B; color: #1E3A8A; border: none;
    padding: 12px 34px; border-radius: 30px; font-size: 15px; font-weight: 600;
    cursor: pointer; transition: all 0.3s; position: relative;
    box-shadow: 0 4px 15px rgba(0,0,0,0.15);
}
.gov-hero .hero-btn:hover { transform: translateY(-2px); box-shadow: 0 6px 24px rgba(0,0,0,0.25); }

/* ===== SECTION TITLE ===== */
.gov-section-title {
    font-size: 20px; font-weight: 700; color: #1E3A8A;
    margin-bottom: 20px; padding-bottom: 12px;
    border-bottom: 3px solid #F59E0B;
}

/* ===== STAT CARDS ===== */
.stat-card {
    background: #FFFFFF; padding: 24px 20px; border-radius: 12px; text-align: center;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06); border: 1px solid #e8ecf1;
    transition: all 0.25s; height: 100%; position: relative; overflow: hidden;
}
.stat-card::before {
    content: ''; position: absolute; top: 0; left: 0; right: 0;
    height: 3px; background: linear-gradient(90deg, #1E3A8A, #2563EB);
}
.stat-card:hover {
    transform: translateY(-4px); box-shadow: 0 8px 24px rgba(0,0,0,0.1);
    border-color: #b8c4d4;
}
.stat-card h3 { font-size: 2.4em; color: #1E3A8A; margin: 0; font-weight: 700; }
.stat-card p { color: #64748B; margin-top: 6px; font-size: 12px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; }
.stat-card .stat-icon { font-size: 30px; margin-bottom: 8px; display: block; }

/* ===== STATUS BADGE ===== */
.status-badge {
    padding: 4px 14px; border-radius: 20px; font-size: 12px;
    font-weight: 600; color: white; display: inline-block;
}

/* ===== DETAIL SECTION ===== */
.detail-section {
    background: #FFFFFF; border-radius: 12px; padding: 26px;
    margin-bottom: 18px; border: 1px solid #e8ecf1;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
}
.detail-section h3 {
    font-size: 16px; font-weight: 600; color: #1E3A8A;
    margin: 0 0 18px 0; padding-bottom: 12px;
    border-bottom: 2px solid #f0f2f5;
}
.detail-label { font-size: 11px; color: #64748B; font-weight: 600; margin-bottom: 3px; text-transform: uppercase; letter-spacing: 0.4px; }
.detail-value { font-size: 15px; color: #0F172A; margin-bottom: 16px; line-height: 1.6; }

/* ===== STEPS ===== */
.step-card {
    background: #FFFFFF; border-radius: 12px; padding: 24px 20px;
    text-align: center; border: 1px solid #e8ecf1; height: 100%;
    transition: all 0.25s;
}
.step-card:hover {
    transform: translateY(-3px); box-shadow: 0 6px 16px rgba(0,0,0,0.07);
}
.step-card .step-number {
    width: 40px; height: 40px;
    background: linear-gradient(135deg, #1E3A8A, #2563EB);
    color: white; border-radius: 50%; display: flex;
    align-items: center; justify-content: center; font-weight: 700;
    font-size: 16px; margin: 0 auto 10px;
}
.step-card h4 { font-size: 14px; font-weight: 600; color: #1E3A8A; margin: 0 0 8px 0; }
.step-card p { font-size: 13px; color: #64748B; line-height: 1.5; margin: 0; }

/* ===== MUROJAAT CARD ===== */
.murojaat-card {
    background: #FFFFFF; border-radius: 12px; padding: 18px 22px;
    margin-bottom: 10px; border: 1px solid #e8ecf1;
    transition: all 0.2s;
}
.murojaat-card:hover {
    border-color: #c0c8d4; box-shadow: 0 2px 10px rgba(0,0,0,0.06);
}
.murojaat-card .title {
    font-size: 15px; font-weight: 600; color: #1E3A8A; margin-bottom: 4px;
}
.murojaat-card .meta { font-size: 13px; color: #64748B; }
.murojaat-card .preview { font-size: 13px; color: #334155; margin-top: 6px; line-height: 1.5; }

/* ===== REESTR ===== */
.reestr-card {
    background: #FFFFFF; border-radius: 12px; padding: 16px 20px;
    margin-bottom: 8px; border: 1px solid #e8ecf1;
    transition: all 0.2s;
}
.reestr-card:hover { border-color: #c8d0db; box-shadow: 0 2px 8px rgba(0,0,0,0.05); }

/* ===== FOOTER ===== */
.gov-footer {
    background: #1E3A8A;
    color: rgba(255,255,255,0.8);
    padding: 35px 0;
    margin-top: 50px;
    width: 100%;
}
.gov-footer .footer-inner {
    max-width: 1400px; margin: 0 auto; padding: 0 40px; text-align: center;
}
.gov-footer p { font-size: 13px; margin: 4px 0; line-height: 1.6; }
.gov-footer .footer-links { display: flex; justify-content: center; gap: 24px; margin-top: 14px; flex-wrap: wrap; }
.gov-footer .footer-links a { color: #F59E0B; text-decoration: none; font-size: 13px; font-weight: 500; }
.gov-footer .footer-links a:hover { text-decoration: underline; }

/* ===== FORM ===== */
.form-section {
    background: #FFFFFF; border-radius: 12px; padding: 26px; margin-bottom: 18px;
    border: 1px solid #e8ecf1; box-shadow: 0 1px 4px rgba(0,0,0,0.04);
}
.form-section h3 { font-size: 16px; font-weight: 600; color: #1E3A8A; margin: 0 0 18px 0; padding-bottom: 10px; border-bottom: 2px solid #f0f2f5; }

/* FILTER SECTION */
.filter-section {
    background: #FFFFFF; border-radius: 12px; padding: 20px 24px;
    margin-bottom: 20px; border: 1px solid #e8ecf1;
}

/* ===== BUTTON STYLES FOR STREAMLIT (asosiy nav va formalar) ===== */
div.stButton > button {
    border-radius: 8px !important;
    font-weight: 500 !important;
    padding: 8px 22px !important;
    transition: all 0.2s !important;
    font-size: 13px !important;
}
div.stButton > button[kind="primary"] {
    background: #1E3A8A !important;
    color: white !important;
    border: none !important;
}
div.stButton > button[kind="primary"]:hover {
    background: #1D4ED8 !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 14px rgba(30,58,138,0.3) !important;
}
div.stButton > button:not([kind="primary"]) {
    background: #FFFFFF !important;
    color: #1E3A8A !important;
    border: 1.5px solid #1E3A8A !important;
}
div.stButton > button:not([kind="primary"]):hover {
    background: #EFF6FF !important;
    border-color: #1E3A8A !important;
}
.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stSelectbox > div > div > select {
    border-radius: 8px !important;
    border: 1.5px solid #d1d9e6 !important;
    font-size: 14px !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus,
.stSelectbox > div > div > select:focus {
    border-color: #2563EB !important;
    box-shadow: 0 0 0 3px rgba(37,99,235,0.1) !important;
}
.stAlert { border-radius: 8px !important; font-size: 14px !important; }
label { font-size: 14px !important; font-weight: 500 !important; }
</style>
"""


def apply_custom_css():
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)


def status_badge(status, status_nomlari=None):
    if status_nomlari is None:
        status_nomlari = get_status_nomlari()
    rang = STATUS_RANGLAR.get(status, "#666")
    nomi = status_nomlari.get(status, status)
    return f'<span class="status-badge" style="background:{rang}">{nomi}</span>'


# ============================================================
# MUROJAATNI WORD (.docx) FORMATIDA SHAKLLANTIRISH
# ============================================================
def generate_murojaat_docx(m, status_nomlari):
    """Bitta murojaat bo'yicha rasmiy shakldagi Word hujjatini (bytes) qaytaradi."""
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), 'Times New Roman')

    for line in ["O'zbekiston Respublikasi Adliya vazirligi",
                 "Fuqarolar tashabbusi platformasi orqali",
                 "kelib tushgan murojaat"]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(13)

    doc.add_paragraph()

    status_nomi = status_nomlari.get(m["status"], m["status"])

    rows = [
        ("1", "Murojaat tartib raqami", f"№ {m['id']}"),
        ("2", "Tushgan sana va vaqt", m["yaratilgan"]),
        ("3", "Murojaatchi F.I.O.", m["fuqaro_ism"]),
        ("4", "Telefon raqami", m["telefon"]),
        ("5", "Elektron manzili", m["email"] or "Ko'rsatilmagan"),
        ("6", "Kategoriya", m["kategoriya"]),
        ("7", "Qonun/tartib asosi", m["qonun_asosi"] or "Ko'rsatilmagan"),
        ("8", "Murojaatning qisqacha mazmuni", m["muammo"]),
        ("9", "Nima uchun xato/adolatsiz deb hisoblaydi", m["nima_uchun_xato"]),
        ("10", "Murojaatchining taklifi", m["taklif"]),
        ("11", "Ko'rib chiqish natijasi (holati)", status_nomi),
        ("12", "Xodim tomonidan berilgan javob", m["javob"] or "—"),
    ]

    table = doc.add_table(rows=len(rows), cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    widths = (Cm(1.0), Cm(6.0), Cm(9.5))

    for row_idx, (num, label, value) in enumerate(rows):
        cells = table.rows[row_idx].cells
        cells[0].text = num
        cells[1].text = label
        cells[2].text = str(value)
        for cell, width in zip(cells, widths):
            cell.width = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for r in para.runs:
                    r.font.size = Pt(11)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cells[1].paragraphs[0].runs:
            cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run(f"Hujjat shakllantirilgan sana: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# HEADER (gerb + sarlavha + til tugmalari, bitta to'q ko'k panelda)
# ============================================================
ADLIYA_GERB_URL = "https://new.adliya.uz/_nuxt/img/logo.78165fc.png"


def _split_title(text):
    """Sarlavhani ikki qatorga bo'lib beradi (rasmdagidek)."""
    words = text.split()
    mid = (len(words) + 1) // 2
    return " ".join(words[:mid]), " ".join(words[mid:])


def render_gov_header():
    line1, line2 = _split_title(t("header_title"))

    # 1-qism: gerb + sarlavha (to'q ko'k fon)
    st.markdown(f"""
    <div class="gov-header">
        <div class="header-inner">
            <div class="header-left">
                <img class="gerb" src="{ADLIYA_GERB_URL}" alt="Gerb"/>
                <div class="title-block">
                    <h1>{line1}<br/>{line2}</h1>
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # 2-qism: til tugmalari — xuddi shu to'q ko'k fonda, sarlavhaning o'ng tomonida
    st.markdown('<div class="page-container header-lang-row" '
                'style="display:flex; justify-content:flex-end; padding-top:4px; padding-bottom:12px;">',
                unsafe_allow_html=True)
    _, c1, c2, c3 = st.columns([12, 1, 1, 1])
    with c1:
        if st.button("UZ", key="lang_uz",
                     type="primary" if st.session_state.til == "uz" else "secondary",
                     use_container_width=True):
            st.session_state.til = "uz"
            st.rerun()
    with c2:
        if st.button("RU", key="lang_ru",
                     type="primary" if st.session_state.til == "ru" else "secondary",
                     use_container_width=True):
            st.session_state.til = "ru"
            st.rerun()
    with c3:
        if st.button("EN", key="lang_en",
                     type="primary" if st.session_state.til == "en" else "secondary",
                     use_container_width=True):
            st.session_state.til = "en"
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)


# ============================================================
# NAVIGATION
# Fayllar orasida (Fuqaro.py <-> pages/1_Adliya_paneli.py) o'tish uchun
# har bir "sahifa" kaliti qaysi faylga tegishli ekanligi belgilangan.
# ============================================================
PAGE_TARGETS = {
    "Bosh sahifa": "Fuqaro.py",
    "Taklif yuborish": "Fuqaro.py",
    "Ochiq reestr": "Fuqaro.py",
    "Ochiq reestr: Batafsil": "Fuqaro.py",
    "Xodim kirishi": "Fuqaro.py",
    "Xodim: Murojaatlar": "pages/1_Adliya_paneli.py",
    "Xodim: Batafsil": "pages/1_Adliya_paneli.py",
    "Xodim: Statistika": "pages/1_Adliya_paneli.py",
}


def render_gov_nav():
    if st.session_state.xodim_id:
        nav_items = [
            ("Bosh sahifa", t("nav_home")),
            ("Ochiq reestr", t("nav_reestr")),
            ("Xodim: Murojaatlar", t("nav_appeals")),
            ("Xodim: Statistika", t("nav_stats")),
            ("__chiqish__", f"{t('nav_logout')} ({st.session_state.xodim_ism})"),
        ]
    else:
        nav_items = [
            ("Bosh sahifa", t("nav_home")),
            ("Taklif yuborish", t("nav_submit")),
            ("Ochiq reestr", t("nav_reestr")),
            ("Xodim kirishi", t("nav_login")),
        ]

    st.markdown('<div class="page-container">', unsafe_allow_html=True)
    actual_cols = st.columns(len(nav_items))

    for col, (key, label) in zip(actual_cols, nav_items):
        with col:
            if key == "__chiqish__":
                if st.button(label, use_container_width=True):
                    st.session_state.xodim_id = None
                    st.session_state.xodim_ism = None
                    st.session_state.sahifa = "Bosh sahifa"
                    st.switch_page("Fuqaro.py")
            else:
                btn_type = "primary" if st.session_state.sahifa == key else "secondary"
                if st.button(label, key=f"nav_{key}", type=btn_type, use_container_width=True):
                    st.session_state.sahifa = key
                    target = PAGE_TARGETS[key]
                    st.switch_page(target)
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown("---")


# ============================================================
# FOOTER
# ============================================================
def render_gov_footer():
    st.markdown(f"""
    <div class="gov-footer">
        <div class="footer-inner">
            <p><strong>{t('header_title')}</strong></p>
            <p>{t('footer_address')}</p>
            <p>{t('footer_contact')}</p>
            <div class="footer-links">
                <a href="#">{t('footer_about')}</a>
                <a href="#">{t('footer_terms')}</a>
                <a href="#">{t('footer_privacy')}</a>
                <a href="#">{t('footer_contact_link')}</a>
                <a href="#">{t('footer_faq')}</a>
            </div>
            <p style="margin-top:16px;">© 2026 {t('footer_copyright')}</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
